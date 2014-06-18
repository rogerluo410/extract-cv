#file name : extract.py
#created at : 2014.06.17
#functionality : read pdf/docx/doc/txt file to text, and extract keyword 'name','phone number','email' from text.

from PyPDF2 import PdfFileWriter, PdfFileReader
from docx import Document
import sys
import math
import re
import os

class ExtractPdf :

   def __init__(self,filename):
      self.filename = filename

   def extract_text_from_pdf(self):
      try:
        input1 = PdfFileReader(open(self.filename, "rb"))
        text = ''
        for iter in range (0, input1.getNumPages()):
            text = text + input1.getPage(iter).extractText().encode("ascii", "ignore")
        self.text = text
        if self.text and not self.text.isspace():
          return [0,"Read text ok"]
        else:
          return [-1,"Read empty text"]
      except IOError:
         return [1,"Cann't read file : "+self.filename]
      except:
         return [1,"Unknown exception"]

   def extract_text_from_pdf1(self):
      try:
         txtfile = 'Temp_pdf.txt'
         rs = os.system('pdftotext '+self.filename+' '+txtfile )
         if rs == 0:
            f_input = open(txtfile)
            line_input = ''
            text =''
            while True :
              line_input = f_input.readline()
              if len(line_input) == 0:
                break
              text = text + line_input
            f_input.close()
            self.text = text
            return [0,"Read text ok"]
         else:
            return [-1,"Failed to convert txt with tool pdftotext"]
      except IOError:
         return [1,"Cann't read file : "+self.filename]
      except:
         return [1,"Unknown exception"]

   def get_all_text(self):
     return self.text


class ExtractDocx:

    def __init__(self,filename):
        self.filename = filename

    def extract_text_from_docx(self):
       try:
        document = Document(self.filename)
        text_chunks = []
        text = ''
        for paragraph in document.paragraphs:
          text_chunks.append(paragraph.text.encode("ascii"))
        for table in document.tables:
            for row in table.rows:
             for cell in row.cells:
               for paragraph in cell.paragraphs:
                  text_chunks.append(paragraph.text.encode("ascii", "ignore"))
        text = text.join(text_chunks[0:])
        self.text = text
        if self.text and not self.text.isspace():
          return [0,"Read read ok"]
        else:
          return [-1,"Read empty text"]
       except IOError:
        return [1,"Can't read file : "+ self.filename ]
       except :
        return [1,"Unknown exception"]

    def extract_text_from_docx1(self):
      try:
         txtfile = 'Temp_docx.txt'
         rs = os.system('docx2txt '+self.filename+' '+txtfile )
         if rs == 0:
            f_input = open(txtfile)
            line_input = ''
            text =''
            while True :
              line_input = f_input.readline()
              if len(line_input) == 0:
                break
              text = text + line_input
            f_input.close()
            self.text = text
            return [0,"Read text ok"]
         else:
            return [-1,"Failed to convert txt with tool pdftotext"]
      except IOError:
         return [1,"Cann't read file : "+self.filename]
      except:
         return [1,"Unknown exception"]

    def get_all_text(self):
       return self.text


class ExtractTxt:
    def __init__(self,filename):
        self.filename = filename

    def extract_text_from_txt(self):
       try:
         f_input = open(self.filename)
         line_input = ''
         text =''
         while True :
          line_input = f_input.readline()
          if len(line_input) == 0:
            break
          text = text + line_input
         f_input.close()
         self.text = text
         if self.text and not self.text.isspace():
           return [0,"Read text ok"]
         else:
           return [-1,"Read empty text"]
       except IOError:
        return [1,"Can't read file : "+ self.filename ]
       except :
        return [1,"Unknown exception"]

    def get_all_text(self):
       return self.text

class ExtractDoc:
    def __init__(self,filename):
        self.filename = filename 
  
    def extract_text_from_doc(self):
      try:
         txtfile = 'Temp_doc.txt'
         rs = os.system('antiword '+self.filename+' >  '+txtfile )
         if rs == 0:
            f_input = open(txtfile)
            line_input = ''
            text =''
            while True :
              line_input = f_input.readline()
              if len(line_input) == 0:
                break
              text = text + line_input
            f_input.close()
            self.text = text
            return [0,"Read text ok"]
         else:
            return [-1,"Failed to convert txt with tool pdftotext"]
      except IOError:
         return [1,"Cann't read file : "+self.filename]
      except:
         return [1,"Unknown exception"]

    def get_all_text(self):
       return self.text
